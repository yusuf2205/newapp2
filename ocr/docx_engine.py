"""Извлечение данных из .docx — текст уже структурирован, OCR не нужен.

Бланк заявки на оплату оформлен таблицей: строка заголовков + строка(и) данных.
Читаем его по колонкам (заголовок -> позиция -> значение той же позиции в строке
данных) — надёжнее, чем искать метку в склеенном тексте таблицы, где после метки
до конца строки идут уже все остальные ячейки. Для документов без такой таблицы
(обычные счета/договоры) используются общие эвристики parse_plain_text.
"""

from __future__ import annotations

import asyncio
import io

from .base import OcrError, parse_plain_text
from .table_utils import goods_from_items_grid, table_fields_from_grid

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _load_document(content: bytes):
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise OcrError("Установите python-docx") from exc

    try:
        return docx.Document(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise OcrError(f"Не удалось прочитать .docx: {exc}") from exc


def _flatten_text(document) -> str:
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _table_grid(table) -> list[list[str]]:
    return [[cell.text for cell in row.cells] for row in table.rows]


def _extract_goods_from_items_table(document) -> str | None:
    for table in document.tables:
        goods = goods_from_items_grid(_table_grid(table))
        if goods:
            return goods
    return None


def _extract_table_fields(document) -> dict:
    for table in document.tables:
        values = table_fields_from_grid(_table_grid(table))
        if values:
            return values
    return {}


async def recognize_docx(content: bytes) -> dict:
    document = await asyncio.to_thread(_load_document, content)
    text = await asyncio.to_thread(_flatten_text, document)
    if not text.strip():
        raise OcrError("В документе не найдено текста")

    result = parse_plain_text(text)
    result.update(await asyncio.to_thread(_extract_table_fields, document))

    if not result.get("goods"):
        result["goods"] = await asyncio.to_thread(_extract_goods_from_items_table, document)

    return result
